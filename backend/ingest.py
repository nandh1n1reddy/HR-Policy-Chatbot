from __future__ import annotations

from pathlib import Path

from backend.config import DOCUMENTS_PATH
from backend.document import Document


SUPPORTED_EXTENSIONS = {".md", ".markdown", ".txt", ".docx", ".pdf"}


def _read_text_file(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(file_path: Path) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(str(file_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs)


def _read_pdf(file_path: Path) -> str:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"# Page {index}\n{text}")
    return "\n\n".join(pages)


def _load_document(file_path: Path, base_path: Path) -> Document | None:
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return None

    if suffix in {".md", ".markdown", ".txt"}:
        content = _read_text_file(file_path)
    elif suffix == ".docx":
        content = _read_docx(file_path)
    elif suffix == ".pdf":
        content = _read_pdf(file_path)
    else:
        content = ""

    if not content.strip():
        return None

    relative = file_path.relative_to(base_path)
    category = relative.parts[0] if len(relative.parts) > 1 else file_path.parent.name

    return Document(
        filename=file_path.name,
        category=category,
        content=content,
        source_path=str(file_path),
    )


def load_documents(source_path: Path | None = None) -> list[Document]:
    base_path = Path(source_path or DOCUMENTS_PATH)
    if not base_path.exists():
        return []

    documents: list[Document] = []
    for file_path in base_path.rglob("*"):
        if file_path.is_file():
            document = _load_document(file_path, base_path)
            if document is not None:
                documents.append(document)
    return documents


if __name__ == "__main__":
    docs = load_documents()
    print(f"Loaded {len(docs)} documents.\n")
    for doc in docs[:5]:
        print("=" * 40)
        print(doc.filename)
        print(doc.category)
        print(len(doc.content))
